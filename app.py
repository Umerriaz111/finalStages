from flask import Flask, render_template, jsonify, request, redirect, url_for, session, flash,send_file, jsonify
import hashlib
import sqlite3
import requests
from bs4 import BeautifulSoup
# from models.Helper import DataBase
from models.supabaseHelper import DataBase
from docx import Document
import os
from huggingface_hub import InferenceClient
from werkzeug.utils import secure_filename
from io import BytesIO  # Make sure to import BytesIO
# from langchain_community.chat_models import ChatOllama
from openai import OpenAI
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import PyPDF2
import docx
import docx2txt
import fitz



# Model ID for Mistral
model_id = "mistralai/Mistral-7B-Instruct-v0.3"
api_token = "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"
# Initialize the client with the model ID and your token
client = InferenceClient(model=model_id, token=api_token)
api_key="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"
app = Flask(__name__)
app.secret_key = 'FullStackProject'
Model=OpenAI(api_key=api_key)

class Database:
    def __init__(self):
        # self.d1 = DataBase()
        self.d1=DataBase()
def documentnameGeneration(input_queries):
    # Prepare the prompt
    prompt = f"""
    You are a knowledgeable assistant. Based on this information: {input_queries}, generate a concise and appropriate heading (maximum 6 words) for a document.
    """
    
    # Generate the response from the model
    response = client.text_generation(prompt, max_new_tokens=20, temperature=0.7, top_p=0.9)
    
    # Extract the heading (assuming response is a string)
    heading = response.strip()  # Remove extra spaces or newlines
    
    return heading

def scrape_link_data(link):
    """Fetch and clean text content from the provided link."""
    try:
        # Get the content of the webpage
        response = requests.get(link)
        response.raise_for_status()

        # Parse the page using BeautifulSoup
        soup = BeautifulSoup(response.content, "html.parser")
        paragraphs = soup.find_all('p')
        text = "\n".join([para.get_text() for para in paragraphs])

        # Return the cleaned text data
        return text

    except requests.exceptions.RequestException as e:
        return f"Error fetching the data from the link: {e}"   

def getresponses(questions,data):
    completion=Model.chat.completions.create(
    model='gpt-4o-mini',
        messages=[
        { "role": "system", "content": f"You are a helpful assistant give the Response based on {data}" },
        {"role":"user","content":questions}
        ],
        max_tokens=4000
    )
    return completion.choices[0].message.content
    

def ask_question_about_link( user_query=None):
    """Generate a response based on the link content and user query."""
    link="https://en.wikipedia.org/wiki/Education_in_India"
    link_data = scrape_link_data(link)

    if "Error" in link_data:
        return link_data


    # Get the response from the LLM
    response = getresponses(user_query,link_data)
    
    return response 

@app.route("/")
def index():
    return render_template("/index.html")

@app.route("/Login")
def Login():
    return render_template("/Login/signin.html")

@app.route("/Signin")
def Signin():
    return render_template("/Login/signup.html")

@app.route('/signup2', methods=['POST'])
def signup2():
    # c1 = Database()
    c1=Database()
    name = request.form['Enter_name_for_signup']
    email = request.form['Enter_email_for_signup']
    password = request.form['Enter_password_for_signup']
    country = request.form['country']
    phone_number = request.form['Phonenumber']
    hashed_password = hashlib.sha256(password.encode()).hexdigest()

    try:
        existing_user = c1.d1.alreadyExistuser(email)
        pn = c1.d1.alreadyExistuserbyphonenumber(phone_number)

        if existing_user and pn:  # Ensure neither is None
            if existing_user['email'] == email and pn['email'] == email:
                flash('Account already exists!', 'error')
                return redirect(url_for('Signin'))
        elif existing_user is None and pn and pn['phonenumber'] == phone_number:  # Check if pn is not None
            flash('Phone Number Attached with some Other Email!', 'error')
            return redirect(url_for('Signin'))
        elif existing_user and existing_user['email'] == email and pn is None:
            flash('Email Attached With Some Other Phone Number', 'error')
            return redirect(url_for('Signin'))

        
        else:
            c1.d1.insert_data(name, email, password, phone_number, country)
            session['u_id'] = c1.d1.get_u_id(email)
            return render_template("/AddTemplates/index.html",name=name)
    except sqlite3.IntegrityError:
        flash('Error creating account!', 'error')
        return redirect(url_for('Signin'))

@app.route('/auth', methods=['POST'])
def auth():
    email = request.form['User_Email']
    password = request.form['User_Password']
    
    c1 = DataBase()
    user = c1.alreadyExistuser(email)
    
    if user is None:
        flash('Invalid Email or Password. User not found.', 'error')
        return render_template("/Login/signin.html")
    
    session['u_id'] = c1.get_u_id(email)
    
    if user['password'] == password:
        print("!!!!!!!!!!")
        print(session['u_id'])
        print("!!!!!!!!!!!!!!!")
        return render_template("/AddTemplates/index.html",name=user['name'])
    else:
        flash('Invalid Email or Password. Try Again!', 'error')
        return render_template("/Login/signin.html")

@app.route('/addtemplate')
def addtemplatepage():
    return render_template('/AddTemplates/addtemplatepage.html')

@app.route('/addtemplates', methods=['POST'])
def processtemplate():
    form_data = request.form.to_dict()
    c1 = Database()
    
    c1.d1.insert_private_template(
        form_data['Templatename'], form_data['PageSize'], form_data['Orientation'], 
        form_data['Margins'], form_data['Columns'], form_data['Font-style'], form_data['Font-size'], 
        form_data['Font-color'], form_data['Font-Styling'], form_data['Text-Alin'], 
        form_data['Line-Spacing'], form_data['Para-Spacing'], form_data['Page-Number'], 
        form_data['Custom-Header'], form_data['Indentation'], form_data['Tab-Stops'], 
        form_data['Bullet-Points'], form_data['Numbered-Lists'], form_data['Table-Borders'], 
        form_data['RowsandColumns'], form_data['Text-Alignment'], form_data['Table_of_content'], 
        session['u_id']
    )
    
    return redirect(url_for("publicandprivatetemplate"))

@app.route('/Publicandprivatetemplate')  
def publicandprivatetemplate():
    c1 = Database()
    privatetemplate = c1.d1.get_Private_template(session['u_id'])
    publictemplate = c1.d1.get_Public_template()
    print("#####################")
    print(privatetemplate)
    print("####################")

    if not publictemplate and not privatetemplate:
        return render_template('/RenderTemplateHistory/index.html')
    elif not publictemplate:
        return render_template('/RenderTemplateHistory/index.html', templates=privatetemplate)
    elif not privatetemplate:
        return render_template('/RenderTemplateHistory/index.html', templates2=publictemplate)
    else:
        return render_template('/RenderTemplateHistory/index.html', templates=privatetemplate, templates2=publictemplate)
            
@app.route('/PublicandprivatetemplateforQuestionAnswers')  
def publicandprivatetemplateforQuestionAnswers():
    c1 = Database()
    privatetemplate = c1.d1.get_Private_template(session['u_id'])
    publictemplate = c1.d1.get_Public_template()

    if not publictemplate and not privatetemplate:
        return render_template('/RenderTemplateHistory/index2.html')
    elif not publictemplate:
        return render_template('/RenderTemplateHistory/index2.html', templates=privatetemplate)
    elif not privatetemplate:
        return render_template('/RenderTemplateHistory/index2.html', templates2=publictemplate)
    else:
        return render_template('/RenderTemplateHistory/index2.html', templates=privatetemplate, templates2=publictemplate)

@app.route('/showinpublic', methods=['POST'])
def showinpublic():
    template_data = request.form.to_dict()
    
    c1 = Database()
    c1.d1.insert_public_template(
        template_data["Templatename"], template_data["PageSize"], template_data["Orientation"], 
        template_data["Margins"], template_data["Columns"], template_data["Font-style"], 
        template_data["Font-size"], template_data["Font-color"], template_data["Font-Styling"], 
        template_data["Text-Alin"], template_data["Line-Spacing"], template_data["Para-Spacing"], 
        template_data["Page-Number"], template_data["Custom-Header"], template_data["Indentation"], 
        template_data["Tab-Stops"], template_data["Bullet-Points"], template_data["Numbered-Lists"], 
        template_data["Table-Borders"], template_data["RowsandColumns"], template_data["Text-Alignment"], 
        template_data["Table_of_content"], session['u_id']
    )
    
    return redirect(url_for("publicandprivatetemplate"))

@app.route('/make_public', methods=['POST'])
def make_public():
    template_data = request.form.to_dict()
    Template_id = int(template_data["Templateid"])
    
    c1 = Database()
    c1.d1.insert_public_template(
        template_data["Templatename"], template_data["PageSize"], template_data["Orientation"], 
        template_data["Margins"], template_data["Columns"], template_data["Font-style"], 
        template_data["Font-size"], template_data["Font-color"], template_data["Font-Styling"], 
        template_data["Text-Alin"], template_data["Line-Spacing"], template_data["Para-Spacing"], 
        template_data["Page-Number"], template_data["Custom-Header"], template_data["Indentation"], 
        template_data["Tab-Stops"], template_data["Bullet-Points"], template_data["Numbered-Lists"], 
        template_data["Table-Borders"], template_data["RowsandColumns"], template_data["Text-Alignment"], 
        template_data["Table_of_content"], session['u_id']
    )
    c1 = Database()
    c1.d1.delete_private_template(Template_id)
    print("###################")
    print(template_data)
    print("###################")
    return redirect(url_for("publicandprivatetemplate"))

@app.route("/delete_public", methods=["POST"])
def deletetemplate():
    Template_id = int(request.form.get("Templateid"))
    
    c1 = Database()
    c1.d1.delete_private_template(Template_id)
    
    return redirect(url_for("publicandprivatetemplate"))

@app.route("/QuestionAnswers", methods=['POST']) 
def Questionanswer():
    Templatename = request.form.get("Templatename")
    return render_template("/QuestionAnswerModel/index.html", Templatename=Templatename)

@app.route("/DocGenerator", methods=['POST'])
def DocGenerator():
    Templatename = request.form.get("Templatename")
    return render_template("/DocGeneratorModel/index.html", Templatename=Templatename)

@app.route('/generate_report', methods=['POST'])
def generate_report():
    # Get user input and additional parameters
    user_input = request.json.get('user_input', '')
    qa_type = request.json.get('qa_type', '')  # e.g., "paragraph"
    number_of_questions = int(request.json.get('number_of_questions', 0))  # Ensure it's an integer

    # Prepare the prompt based on QA type and number of questions
    prompt = f"""
    You are a highly knowledgeable assistant. Based on the following user query, generate {number_of_questions} relevant questions and answers in {qa_type} format.

    User Query: {user_input}

    Please provide the output in the following format:

    Question: [Generated Question]
    Answer: [Generated Answer]
    """

    # # Initialize the LLM from Ollama
    # local_model = "mistral"
    # llm = ChatOllama(model=local_model)

    # # Invoke the model to get the report
    # response = llm.invoke(prompt)  # Use the modified prompt
    
    # Get model response with parameters for longer responses
    
   # Model completion API call
    completion = Model.chat.completions.create(
    model='gpt-4o-mini',
    messages=[
        { "role": "system", "content": "You are a helpful assistant." },
        {"role": "user", "content": prompt}
    ],
    max_tokens=4000
)

# Output the completion response
    # response = completion['choices'][0]['message']['content']
    # Access the content directly
    response = completion.choices[0].message.content
    # report = response.content  # Update this line based on your AIMessage implementation
    # print(report)
    # documentname=documentnameGeneration(user_input)
    # print("!!!!!!!!!!!!!!!!!!!")
    # print(documentname)
    # print("!!!!!!!!!!!!!!!!!!!")
    # Create a Word document
    doc = Document()
    doc.add_heading('Brief Report', level=1)

    # Add parameters to the report
    doc.add_paragraph(f"QA Type: {qa_type.capitalize()}")
    doc.add_paragraph(f"Number of Questions: {number_of_questions}")
    doc.add_paragraph("")  # Add a blank line for spacing

    # Format the report content
    for entry in response.split("\n\n"):  # Assuming each Q&A is separated by two newlines
        if entry.strip():  # Skip empty entries
            # Extract question and answer from the entry
            if "Answer:" in entry:
                question, answer = entry.split('Answer:', 1)  # Split on the first occurrence of 'Answer:'
                question = question.replace('Question:', '').strip()  # Clean up the question
                answer = answer.strip()  # Clean up the answer
                
                # Add to the Word document
                doc.add_heading(f'Question: {question}', level=2)
                doc.add_paragraph(f'Answer: {answer}')

    # Save the document to a BytesIO object
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)
    
    # Send the document as a downloadable file
    return send_file(
        report_stream,
        as_attachment=True,
        download_name='report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    

@app.route('/generate_report_for_tf', methods=['POST'])
def generate_report_tf():
    # Get user input and additional parameters
    user_input = request.json.get('user_input', '')
    number_of_questions = int(request.json.get('number_of_questions', 0))  # Ensure it's an integer

    # Prepare the prompt based on QA type and number of questions
    prompt = f"""
    You are a highly knowledgeable assistant. Based on the following user query, generate {number_of_questions} questions related to the query with only True/False answers. 
    Do not provide any additional information or explanations, only the question and a True/False answer.

    User Query: {user_input}

    Please provide the output in this strict format:
    Question 1: [Generated Question]
    Answer: True/False

    Example:
    Question 1: Pakistan is a Muslim country?
    Answer: True

    Question 2: America is a Muslim country?
    Answer: False
    """

    # # Initialize the LLM from Ollama
    # local_model = "mistral"
    # llm = ChatOllama(model=local_model)

    # # Invoke the model to get the report
    # response = llm.invoke(prompt)  # Use the modified prompt

    # # Access the content directly
    # report = response.content
    # print(report)
    # print("########################")
    # response = client.text_generation(prompt, max_new_tokens=1500, temperature=0.5, top_p=0.9)
    # print(response)
    # print("########################")
    completion = Model.chat.completions.create(
    model='gpt-4o-mini',
    messages=[
        { "role": "system", "content": "You are a helpful assistant." },
        {"role": "user", "content": prompt}
    ],
    max_tokens=4000
)
    response = completion.choices[0].message.content
    
    # Create a Word document
    doc = Document()
    doc.add_heading('Generated Report', level=1)

    # Add parameters to the report
    doc.add_paragraph(f"User Query: {user_input}")
    doc.add_paragraph(f"Number of Questions: {number_of_questions}")
    doc.add_paragraph("")  # Add a blank line for spacing

    # Format the report content to only include questions and answers
    for entry in response.split("\n\n"):  # Assuming each Q&A is separated by two newlines
        if entry.strip():  # Skip empty entries
            # Extract question and answer from the entry
            lines = entry.splitlines()
            question = lines[0].replace('Question', '').strip()  # Clean up the question
            answer = lines[1].replace('Answer', '').strip()  # Clean up the answer

            # Add the question and answer to the document
            doc.add_heading(f'{question}', level=2)
            doc.add_paragraph(f'Answer: {answer}')

    # Save the document to a BytesIO object
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)

    # Send the document as a downloadable file
    return send_file(
        report_stream,
        as_attachment=True,
        download_name='report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/generate_report_for_fb', methods=['POST'])
def generate_report_fb():
    user_input = request.json.get('user_input', '')
    number_of_questions = int(request.json.get('number_of_questions', 0))

    # Updated prompt ensuring dashes in questions and correct answers
    prompt = f"""
    You are a knowledgeable assistant. Based on the user query, generate {number_of_questions} fill-in-the-blank questions with dashes (__________) in the question and a single correct answer.

    Example:
    Question: India gained its ___________ in 1947.
    Answer: a) Independence

    User Query: {user_input}

    Output format:
    Question: [Generated Question with _______ for blanks]
    Answer: [Correct Answer]
    """

    # # Generate response using the local model (e.g., Mistral, etc.)
    # llm = ChatOllama(model="mistral")
    # response = llm.invoke(prompt)
    # report = response.content
    # print(report)  # This prints the correct response to the console
    # print("##############################")
    # response = client.text_generation(prompt, max_new_tokens=1500, temperature=0.5, top_p=0.9)
    # print(response)
    # print("##############################")
    completion = Model.chat.completions.create(
    model='gpt-4o-mini',
    messages=[
        { "role": "system", "content": "You are a helpful assistant." },
        {"role": "user", "content": prompt}
    ],
    max_tokens=4000
)
    response = completion.choices[0].message.content
    
    # Create a Word document
    doc = Document()
    doc.add_heading('Generated Report', level=1)

    # Add user input and other details
    doc.add_paragraph(f"User Query: {user_input}")
    doc.add_paragraph(f"Number of Questions: {number_of_questions}")
    doc.add_paragraph("")  # Blank line for spacing

    # Process the response for each question and answer
    questions_and_answers = response.strip().split("\n\n")
    for entry in questions_and_answers:
        if entry.strip():
            lines = entry.splitlines()

            # Extract question and answer
            question = lines[0].replace('Question:', '').strip()
            answer = lines[1].replace('Answer:', '').strip() if len(lines) > 1 else 'Answer not available'

            # Add question to the document
            doc.add_heading(f'Question: {question}', level=2)

            # Add answer to the document
            doc.add_paragraph(f"Answer: {answer}")

            # Blank line for spacing between questions
            doc.add_paragraph("")

    # Save the document to an in-memory buffer
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)

    # Return the generated document as a downloadable file
    return send_file(
        report_stream,
        as_attachment=True,
        download_name='report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# Function to extract text from .docx
def extract_text_from_docx(docx_file_path):
    doc = docx.Document(docx_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to extract text from .doc
def extract_text_from_doc(doc_file_path):
    # Use docx2txt to extract text
    text = docx2txt.process(doc_file_path)
    return text

# Function to extract text from .pdf
def extract_text_from_pdf(pdf_file_path):
    pdf_text = ""
    with fitz.open(pdf_file_path) as doc:
        for page in doc:
            pdf_text += page.get_text()
    return pdf_text

@app.route('/generate_report_from_doc', methods=['POST'])
def generate_report_from_doc():
    # Handle file upload
    file = request.files['file']
    filename = secure_filename(file.filename)
    file_ext = filename.split('.')[-1].lower()
    file_path = os.path.join('/Users/umerriaz/Desktop/Frontend/uploadedFiles', filename)
    file.save(file_path)

    # Extract clean text from the uploaded file based on the file type
    if file_ext == 'docx':
        user_input = extract_text_from_docx(file_path)
    elif file_ext == 'doc':
        user_input = extract_text_from_doc(file_path)
    elif file_ext == 'pdf':
        user_input = extract_text_from_pdf(file_path)
    else:
        return {"error": "Unsupported file format"}, 400

    # Remove images (handled inherently by extraction functions)

    # QA Type and Number of Questions
    qa_type = request.form.get('qa_type', '')
    number_of_questions = int(request.form.get('number_of_questions', 0))

    # Prepare the prompt for LLM
    prompt = f"""
    You are a highly knowledgeable assistant. Based on the following user query, generate {number_of_questions} relevant questions and answers in {qa_type} format.

    User Query: {user_input}

    Please provide the output in the following format:

    Question: [Generated Question]
    Answer: [Generated Answer]
    """

    # # Initialize the LLM from Ollama
    # local_model = "mistral"
    # llm = ChatOllama(model=local_model)

    # # Invoke the model to get the report
    # response = llm.invoke(prompt)
    # report = response.content
    print("##############################")
    response = client.text_generation(prompt, max_new_tokens=1500, temperature=0.5, top_p=0.9)
    print(response)
    print("##############################")

    # Create a Word document
    doc = Document()
    doc.add_heading('User Text', level=1)
    doc.add_paragraph(user_input)
    doc.add_heading('Brief Report', level=1)
    
    # Add parameters to the report
    doc.add_paragraph(f"QA Type: {qa_type.capitalize()}")
    doc.add_paragraph(f"Number of Questions: {number_of_questions}")
    doc.add_paragraph("")  # Blank line for spacing

    # Format the report content
    for entry in response.split("\n\n"):  # Assuming each Q&A is separated by two newlines
        if entry.strip():  # Skip empty entries
            if "Answer:" in entry:
                question, answer = entry.split('Answer:', 1)
                question = question.replace('Question:', '').strip()
                answer = answer.strip()

                # Add to the Word document
                doc.add_heading(f'Question: {question}', level=2)
                doc.add_paragraph(f'Answer: {answer}')

    # Save the document to a BytesIO object
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)

    # Delete the file after processing
    if os.path.exists(file_path):
        os.remove(file_path)

    # Send the document as a downloadable file
    return send_file(
        report_stream,
        as_attachment=True,
        download_name='report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    
# True /False Generation from Document 

@app.route('/generate_tf_from_doc', methods=['POST'])
def generate_report_from_doc_for_tf():
    # Handle file upload
    file = request.files['file']
    filename = secure_filename(file.filename)
    file_ext = filename.split('.')[-1].lower()
    file_path = os.path.join('/Users/umerriaz/Desktop/Frontend/uploadedFiles', filename)
    file.save(file_path)

    # Extract clean text from the uploaded file based on the file type
    if file_ext == 'docx':
        user_input = extract_text_from_docx(file_path)
    elif file_ext == 'doc':
        user_input = extract_text_from_doc(file_path)
    elif file_ext == 'pdf':
        user_input = extract_text_from_pdf(file_path)
    else:
        return {"error": "Unsupported file format"}, 400

    # QA Type and Number of Questions
    number_of_questions = int(request.form.get('number_of_questions', 0))

    # Prepare the prompt to ensure the answers are either "True" or "False"
    prompt = f"""
        You are an AI assistant. Based on the following user query, generate {number_of_questions} True/False questions and respond with only "True" or "False." Do not include any additional explanations or qualifiers like "assuming user's statement is accurate." Assume the user's statements are factual, and respond strictly with either "True" or "False."

        User Query: {user_input}

        Please provide the output strictly in this format:
        Question 1: [Generated Question]
        Answer: True or False

        Example:
        Question 1: Pakistan is a Muslim country?
        Answer: True

        Question 2: America is a Muslim country?
        Answer: False
    """



    # # Initialize the LLM from Ollama
    # local_model = "mistral"
    # llm = ChatOllama(model=local_model)

    # # Invoke the model to get the report
    # response = llm.invoke(prompt)

    # # Access the content directly
    # report = response.content
    # print(report)
    print("##############################")
    response = client.text_generation(prompt, max_new_tokens=1500, temperature=0.5, top_p=0.9)
    print(response)
    print("##############################")

    # Create a Word document
    doc = Document()
    doc.add_heading("User Input", level=1)
    doc.add_paragraph(user_input)
    doc.add_paragraph(" ")
    doc.add_heading('Generated Report', level=1)
    doc.add_paragraph("")  # Add a blank line for spacing

    # Format the report content to only include questions and answers
    for entry in response.split("\n\n"):  # Assuming each Q&A is separated by two newlines
        if entry.strip():  # Skip empty entries
            # Extract question and answer from the entry
            lines = entry.splitlines()
            question = lines[0].replace('Question', '').strip()  # Clean up the question
            answer = lines[1].replace('Answer:', '').strip()  # Clean up the answer

            # Add the question and answer to the document
            doc.add_heading(f'{question}', level=2)
            doc.add_paragraph(f'Answer: {answer}')

    # Save the document to a BytesIO object
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)
    
    # Delete the file after processing
    if os.path.exists(file_path):
        os.remove(file_path)

    # Send the document as a downloadable file
    return send_file(
        report_stream,
        as_attachment=True,
        download_name='report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/generate_fb_from_doc', methods=['POST'])
def generate_report_from_doc_for_fb():
    # Handle file upload
    file = request.files['file']
    filename = secure_filename(file.filename)
    file_ext = filename.split('.')[-1].lower()
    file_path = os.path.join('/Users/umerriaz/Desktop/Frontend/uploadedFiles', filename)
    file.save(file_path)

    # Extract clean text from the uploaded file based on the file type
    if file_ext == 'docx':
        user_input = extract_text_from_docx(file_path)
    elif file_ext == 'doc':
        user_input = extract_text_from_doc(file_path)
    elif file_ext == 'pdf':
        user_input = extract_text_from_pdf(file_path)
    else:
        return {"error": "Unsupported file format"}, 400

    # QA Type and Number of Questions
    number_of_questions = int(request.form.get('number_of_questions', 0))

     # Updated prompt ensuring dashes in questions and correct answers
    prompt = f"""
            You are a knowledgeable assistant. Based on the user query, generate {number_of_questions} fill-in-the-blank questions with dashes (__________) in the question and a single correct answer.

            Example:
            Question: India gained its ___________ in 1947.
            Answer: a) Independence

            User Query: {user_input}

            Output format:
            Question: [Generated Question with _______ for blanks]
            Answer: [Correct Answer]
    """



    # # Generate response using the local model (e.g., Mistral, etc.)
    # llm = ChatOllama(model="mistral")
    # response = llm.invoke(prompt)
    # report = response.content
    # print(report)  # This prints the correct response to the console
    print("##############################")
    response = client.text_generation(prompt, max_new_tokens=1500, temperature=0.5, top_p=0.9)
    print(response)
    print("##############################")
    # Create a Word document
    doc = Document()
    doc.add_heading("User Input",level=1)
    doc.add_paragraph(user_input)
    # doc.add_heading('Generated Report', level=1)

    # Add user input and other details
    # doc.add_paragraph(f"User Query: {user_input}")
    doc.add_paragraph(f"Number of Questions: {number_of_questions}")
    doc.add_paragraph("")  # Blank line for spacing

    # Process the response for each question and answer
    questions_and_answers = response.strip().split("\n\n")
    for entry in questions_and_answers:
        if entry.strip():
            lines = entry.splitlines()

            # Extract question and answer
            question = lines[0].replace('Question:', '').strip()
            answer = lines[1].replace('Answer:', '').strip() if len(lines) > 1 else 'Answer not available'

            # Add question to the document
            doc.add_heading(f'Question: {question}', level=2)

            # Add answer to the document
            doc.add_paragraph(f"Answer: {answer}")

            # Blank line for spacing between questions
            doc.add_paragraph("")

    # Save the document to an in-memory buffer
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)
    
    # Delete the file after processing
    if os.path.exists(file_path):
        os.remove(file_path)

    # Return the generated document as a downloadable file
    return send_file(
        report_stream,
        as_attachment=True,
        download_name='report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/documentEditorView')
def documentEditorView():
    return render_template("/DocumentEditor/grid.html")

@app.route('/chatbot')
def chatbot():
    return render_template('chatbot/chatbot3.html')
@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json.get('message')
    # Example response logic
    response=ask_question_about_link(user_input)
    
    return jsonify({'response': response})



if __name__ == '__main__':
    app = app.run(debug=True)
